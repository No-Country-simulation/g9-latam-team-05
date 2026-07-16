import sys
import subprocess
import os

try:
    from docx import Document
except ImportError:
    print("Installing python-docx library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ─── HELPERS ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:color="auto" w:val="clear"/>')
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex=(0x1E, 0x3A, 0x8A), space_before=18, space_after=8):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(*color_hex)
    run.font.name = "Arial"
    return h

def add_paragraph(doc, text, space_after=10, bold_first=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    if bold_first:
        run_b = p.add_run(bold_first)
        run_b.bold = True
        run_b.font.name = "Arial"
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_bullet(doc, text, bold_prefix=None, color=(0x33, 0x41, 0x55)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(15)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = "Arial"
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_table_header_row(table, headers, bg_hex="1E3A8A"):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

def add_table_data_row(table, row_idx, values, bg_hex=None):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        if bg_hex:
            set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

# ─── DOCUMENT BUILDER ────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ── Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════════
    # PORTADA
    # ════════════════════════════════════════════════
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)
    r = cover.add_run("HACKATHON ONE — ALURA + ORACLE\nG9 LATAM — EQUIPO 05")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(12)
    rt = title.add_run("FINANCE AI\nAsistente Inteligente de Salud Financiera")
    rt.font.name = "Arial"
    rt.font.size = Pt(26)
    rt.bold = True
    rt.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(60)
    rs = subtitle.add_run(
        "Documento de Especificación de Requerimientos y Arquitectura del Frontend\n"
        "Pantallas, Funcionalidades, Flujos de Integración y Diseño Visual del Sistema"
    )
    rs.font.name = "Arial"
    rs.font.size = Pt(13)
    rs.italic = True
    rs.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(80)
    rm = meta.add_run("Stack Tecnológico: Angular v17+  |  Spring Boot 3.x  |  Python ML  |  Oracle Cloud Infrastructure (OCI)\nJulio 2026")
    rm.font.name = "Arial"
    rm.font.size = Pt(10)
    rm.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 1. INTRODUCCIÓN Y OBJETIVO
    # ════════════════════════════════════════════════
    add_heading(doc, "1. Introducción y Objetivo del Sistema", level=1)

    add_paragraph(doc,
        "Finance AI es una solución inteligente diseñada para transformar datos financieros brutos en "
        "conocimiento accionable. Inspirada en el modelo de las mejores aplicaciones fintech del mercado "
        "(Fintonic, Money Manager, Splitwise), va más allá al incorporar Inteligencia Artificial para "
        "clasificar gastos automáticamente, evaluar el comportamiento financiero del usuario y generar "
        "recomendaciones personalizadas de forma dinámica."
    )
    add_paragraph(doc,
        "El sistema atiende una necesidad real: las personas tienen acceso a sus transacciones bancarias "
        "pero carecen de las herramientas para interpretar ese volumen de datos. Finance AI resuelve esto "
        "con una interfaz visual intuitiva, gráficos en tiempo real y un motor de análisis predictivo "
        "que clasifica automáticamente cada gasto y genera un diagnóstico del perfil financiero."
    )

    add_heading(doc, "1.1 Objetivos Específicos del Frontend", level=2, space_before=12)
    add_bullet(doc, "Proporcionar una experiencia de usuario premium, intuitiva y completamente responsive.", bold_prefix="UX Premium: ")
    add_bullet(doc, "Permitir al usuario simular, registrar y clasificar sus transacciones de forma manual o masiva.", bold_prefix="Captura de Datos: ")
    add_bullet(doc, "Renderizar en tiempo real indicadores financieros (KPIs) y gráficos dinámicos actualizados por Signals.", bold_prefix="Visualización: ")
    add_bullet(doc, "Consumir los endpoints REST del backend (Spring Boot + Python ML) y presentar resultados de forma visual.", bold_prefix="Integración API: ")
    add_bullet(doc, "Mantener un historial de análisis para rastrear la evolución del comportamiento financiero.", bold_prefix="Historial: ")

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 2. ARQUITECTURA DEL SISTEMA (VISIÓN COMPLETA E2E)
    # ════════════════════════════════════════════════
    add_heading(doc, "2. Arquitectura del Sistema Completo (End-to-End)", level=1)

    add_paragraph(doc,
        "El sistema Finance AI adopta una arquitectura de tres capas desacopladas que colaboran "
        "de forma asíncrona para garantizar escalabilidad, mantenibilidad y bajo acoplamiento entre componentes:"
    )

    add_heading(doc, "Capa 1 — Frontend (Angular v17+)", level=2, color_hex=(0x28, 0x35, 0x93), space_before=10)
    add_paragraph(doc,
        "Es la capa de presentación e interacción. Implementada con Angular moderno usando componentes "
        "Standalone, Signals para reactividad, Formularios Reactivos fuertemente tipados y Chart.js para "
        "gráficos dinámicos. Se comunica exclusivamente con el backend via HTTPS/REST en formato JSON."
    )

    add_heading(doc, "Capa 2 — Backend (Java Spring Boot 3.x)", level=2, color_hex=(0x2E, 0x7D, 0x32), space_before=10)
    add_paragraph(doc,
        "Actúa como el núcleo de negocio, seguridad y orquestación. Expone los endpoints REST documentados "
        "con OpenAPI/Swagger, valida los datos de entrada, aplica las reglas de negocio financiero y "
        "delega el procesamiento predictivo al microservicio de Python vía HTTP interno. "
        "Persiste el historial de análisis en la base de datos de OCI."
    )

    add_heading(doc, "Capa 3 — Microservicio de IA (Python + Scikit-Learn, desplegado en OCI)", level=2, color_hex=(0x15, 0x65, 0xC0), space_before=10)
    add_paragraph(doc,
        "Contiene los modelos de Machine Learning (clasificación supervisada) entrenados con datos "
        "financieros del equipo. Recibe las descripciones textuales de transacciones y devuelve la "
        "categoría predicha (Alimentación, Transporte, Salud, etc.) y el perfil financiero del usuario "
        "(Saludable, En Observación, En Riesgo) junto con su probabilidad estadística."
    )

    add_heading(doc, "OCI (Oracle Cloud Infrastructure)", level=2, color_hex=(0xE6, 0x51, 0x00), space_before=10)
    add_bullet(doc, "Object Storage: almacena los modelos ML serializados (.joblib/.pkl).")
    add_bullet(doc, "OCI Compute / Functions: hospeda el microservicio de Python.")
    add_bullet(doc, "OCI Database (Opcional): persiste el historial de análisis de cada usuario.")

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 3. MAPA COMPLETO DE PANTALLAS DEL SISTEMA
    # ════════════════════════════════════════════════
    add_heading(doc, "3. Mapa Completo de Pantallas del Sistema Frontend", level=1)

    add_paragraph(doc,
        "A continuación se detalla el mapa completo de rutas y vistas de la aplicación Finance AI en Angular. "
        "Cada pantalla corresponde a una Feature con lazy loading independiente, garantizando rendimiento óptimo:"
    )

    # Tabla de pantallas
    screen_table = doc.add_table(rows=12, cols=4)
    screen_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    screen_table.style = "Table Grid"
    add_table_header_row(screen_table, ["#", "Ruta (URL)", "Nombre de Pantalla", "Descripción Breve"])
    screen_data = [
        ("1", "/login",               "Inicio de Sesión",                "Autenticación con email y contraseña. JWT generado por Spring Boot."),
        ("2", "/register",            "Registro de Usuario",             "Formulario de alta: nombre, email, contraseña. Validaciones reactivas en Angular."),
        ("3", "/onboarding",          "Configuración Inicial del Perfil","Ingreso mensual, nivel de endeudamiento (%), frecuencia de ahorro. Solo la primera vez."),
        ("4", "/dashboard",           "Dashboard Principal",             "Panel de KPIs, gráfico de dona de gastos por categoría y semáforo de perfil financiero."),
        ("5", "/transacciones",       "Gestión de Transacciones",        "Tabla interactiva de gastos. Registro manual y carga masiva por CSV. Clasificación automática."),
        ("6", "/simulador",           "Simulador de Escenarios",         "Simula el impacto de nuevos gastos o ahorros sobre el perfil financiero antes de confirmarlos."),
        ("7", "/perfil-financiero",   "Diagnóstico de Perfil",           "Resultado completo: perfil (Saludable/En Riesgo), probabilidad y resumen de gastos por categoría."),
        ("8", "/recomendaciones",     "Recomendaciones Inteligentes",    "Tarjetas con consejos personalizados generados por el modelo de IA según el comportamiento detectado."),
        ("9", "/historial",           "Historial de Evolución",          "Gráfico de líneas/barras de la evolución del comportamiento financiero mes a mes."),
        ("10", "/configuracion",      "Configuración de Cuenta",         "Editar perfil financiero base (ingreso, endeudamiento, ahorro) y datos de la cuenta."),
        ("11", "/exportar",           "Exportación de Informes",         "Genera reporte en PDF o CSV del historial de transacciones y análisis del período seleccionado."),
    ]
    for i, row_data in enumerate(screen_data, start=1):
        bg = "F8FAFC" if i % 2 == 0 else None
        add_table_data_row(screen_table, i, row_data, bg_hex=bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 4. DETALLE DE CADA PANTALLA
    # ════════════════════════════════════════════════
    add_heading(doc, "4. Especificación Detallada de Cada Pantalla", level=1)

    # ── 4.1 Login y Registro
    add_heading(doc, "4.1 Autenticación: Login y Registro (/login, /register)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Las pantallas de autenticación son el punto de entrada del sistema. Implementadas con Formularios "
        "Reactivos fuertemente tipados de Angular, validan en tiempo real cada campo antes de habilitarse "
        "el botón de envío, evitando peticiones innecesarias al backend."
    )
    add_bullet(doc, "Campos de Login: email (validación de formato), contraseña (mínimo 8 caracteres).")
    add_bullet(doc, "Campos de Registro: nombre completo, email (único), contraseña, confirmación de contraseña.")
    add_bullet(doc, "Seguridad: el backend genera un token JWT al autenticarse; el frontend lo almacena en memoria (no localStorage) y lo adjunta en cada petición mediante un interceptor funcional de Angular.")
    add_bullet(doc, "Ruta protegida: un AuthGuard funcional redirige al login si el token expiró o no existe.")
    add_bullet(doc, "Diseño: pantalla dividida — lado izquierdo con branding/frase motivacional de Finance AI, lado derecho con el formulario en modo glassmorphism oscuro.")

    # ── 4.2 Onboarding
    add_heading(doc, "4.2 Onboarding — Configuración Inicial del Perfil (/onboarding)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Se muestra una única vez tras el registro. Recolecta la información financiera base del usuario "
        "que será usada como contexto permanente en todos los análisis. Un indicador de progreso visual "
        "(stepper de 3 pasos) guía al usuario sin abrumarlo con todos los campos a la vez."
    )
    add_bullet(doc, "Paso 1 — ¿Cuánto ganas?: Ingreso mensual neto en moneda local.")
    add_bullet(doc, "Paso 2 — ¿Cuánto debes?: Nivel de endeudamiento como porcentaje del ingreso (slider 0-100%).")
    add_bullet(doc, "Paso 3 — ¿Con qué frecuencia ahorras?: Selector visual (Baja / Media / Alta).")
    add_bullet(doc, "Al finalizar: se envían los datos al backend y se redirige al Dashboard con animación de bienvenida.")

    # ── 4.3 Dashboard
    add_heading(doc, "4.3 Dashboard Principal (/dashboard)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Es la pantalla más importante del sistema: el panel de control en tiempo real. Actualiza "
        "todos sus componentes automáticamente cuando el estado global (Signals) cambia."
    )
    add_bullet(doc, "KPIs superiores (4 tarjetas): Ingreso Mensual, Total de Gastos del Período, Saldo Disponible, Tasa de Ahorro (%).")
    add_bullet(doc, "Semáforo de Perfil Financiero: banner adaptativo con color verde/ámbar/rojo y el nombre del perfil (Saludable, En Observación, En Riesgo).")
    add_bullet(doc, "Gráfico de Dona (Chart.js): distribución porcentual de gastos por categoría del período actual.")
    add_bullet(doc, "Gráfico de Barras: comparativa de gastos del mes anterior vs mes actual por categoría.")
    add_bullet(doc, "Accesos rápidos: botones flotantes para 'Agregar Transacción' y 'Ver Recomendaciones'.")

    # ── 4.4 Transacciones
    add_heading(doc, "4.4 Gestión de Transacciones (/transacciones)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Núcleo operativo del sistema. Permite al usuario alimentar el motor de análisis con sus gastos "
        "reales o simulados. Soporta dos métodos de captura y presenta la clasificación automática al instante."
    )
    add_bullet(doc, "Tabla de transacciones: columnas Fecha, Descripción, Monto, Categoría (con icono y badge de color), Acciones (editar/eliminar).")
    add_bullet(doc, "Entrada manual: formulario inline con campo descripción y monto. Al confirmar, envía POST al backend y la fila aparece en la tabla con la categoría ya clasificada.")
    add_bullet(doc, "Carga de CSV (Lotes): componente Drag & Drop. Parsea el archivo en el frontend, muestra una previsualización de las filas detectadas y permite confirmar el envío al backend para clasificación masiva.")
    add_bullet(doc, "Filtros: por categoría, rango de fechas y monto mínimo/máximo.")
    add_bullet(doc, "Ordenación: por fecha o monto (ascendente y descendente).")

    # ── 4.5 Simulador
    add_heading(doc, "4.5 Simulador de Escenarios (/simulador)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Esta pantalla es un diferenciador clave del sistema. Permite al usuario jugar con escenarios "
        "hipotéticos antes de confirmar un gasto o cambio de hábito, viendo en tiempo real cómo impactaría "
        "su perfil financiero."
    )
    add_bullet(doc, "Panel izquierdo: formulario para añadir transacciones simuladas (descripción + monto) o modificar el ingreso/endeudamiento temporalmente.")
    add_bullet(doc, "Panel derecho: previsualización del nuevo perfil calculado (Saludable/En Riesgo) y las recomendaciones que generaría este escenario.")
    add_bullet(doc, "Botón 'Confirmar Escenario': convierte las transacciones simuladas en reales y las registra en el sistema.")
    add_bullet(doc, "Botón 'Descartar': limpia el simulador sin guardar nada.")

    # ── 4.6 Perfil Financiero
    add_heading(doc, "4.6 Diagnóstico de Perfil Financiero (/perfil-financiero)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Pantalla dedicada al resultado completo del análisis generado por el backend. Es la "
        "visualización del JSON de respuesta del endpoint POST /analisis-financiero presentada "
        "de forma clara y atractiva."
    )
    add_bullet(doc, "Gauge (indicador circular): muestra la probabilidad del perfil (ej: 82% 'En Observación').")
    add_bullet(doc, "Resumen de gastos por categoría: tabla con ícono, categoría y monto total del período.")
    add_bullet(doc, "Evolución de perfil: historial de los últimos perfiles obtenidos (ej: el mes pasado era 'En Riesgo', este mes pasó a 'En Observación').")

    # ── 4.7 Recomendaciones
    add_heading(doc, "4.7 Recomendaciones Inteligentes (/recomendaciones)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Presenta de forma visual e interactiva los consejos generados por el modelo de IA en base "
        "al comportamiento financiero detectado."
    )
    add_bullet(doc, "Tarjetas de Recomendación: cada consejo es una tarjeta con ícono, categoría afectada, descripción del problema y acción sugerida concreta.")
    add_bullet(doc, "Prioridad visual: las alertas críticas aparecen en rojo, las de observación en ámbar y las positivas en verde.")
    add_bullet(doc, "Ejemplos generados: 'Tus gastos de Ocio superan el 20% de tu ingreso. Reducir $150/mes mejoraría tu perfil a Saludable.'")

    # ── 4.8 Historial
    add_heading(doc, "4.8 Historial de Evolución Financiera (/historial)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Permite al usuario realizar el seguimiento de la evolución de su comportamiento financiero "
        "a lo largo del tiempo, cumpliendo uno de los requisitos explícitos del Hackathon."
    )
    add_bullet(doc, "Selector de período: mensual, trimestral o anual.")
    add_bullet(doc, "Gráfico de líneas: evolución del saldo disponible, tasa de ahorro y nivel de endeudamiento en el tiempo.")
    add_bullet(doc, "Timeline de perfiles: registro cronológico de los perfiles financieros obtenidos (Saludable/En Observación/En Riesgo).")

    # ── 4.9 Exportar
    add_heading(doc, "4.9 Exportación de Informes (/exportar)", level=2, color_hex=(0x1E, 0x3A, 0x8A))
    add_paragraph(doc,
        "Funcionalidad opcional de alto valor que diferencia al sistema en el Hackathon. "
        "Permite al usuario obtener un respaldo tangible de su historial financiero."
    )
    add_bullet(doc, "Exportar transacciones a CSV: descarga el listado filtrado por período seleccionado.")
    add_bullet(doc, "Exportar informe a PDF: genera un resumen ejecutivo con KPIs, gráficos y recomendaciones del período.")

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 5. CONTRATO DE DATOS Y ENDPOINTS
    # ════════════════════════════════════════════════
    add_heading(doc, "5. Contratos de la API REST (Frontend ↔ Backend)", level=1)

    add_paragraph(doc,
        "El servicio FinanceService de Angular consume los siguientes endpoints documentados en el backend Spring Boot:"
    )

    # Tabla de endpoints
    ep_table = doc.add_table(rows=7, cols=4)
    ep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ep_table.style = "Table Grid"
    add_table_header_row(ep_table, ["Método", "Endpoint", "Descripción", "Usado en Pantalla"])
    ep_data = [
        ("POST", "/api/auth/login",            "Autenticación. Retorna JWT.",                               "Login"),
        ("POST", "/api/auth/register",         "Registro de nuevo usuario.",                                "Register"),
        ("POST", "/api/perfil",                "Guarda/actualiza el perfil financiero base del usuario.",   "Onboarding / Config"),
        ("POST", "/analisis-financiero",       "Clasificación de transacciones + diagnóstico de perfil.",   "Transacciones / Perfil / Simulador"),
        ("GET",  "/api/transacciones",         "Lista historial de transacciones del usuario autenticado.", "Transacciones / Historial"),
        ("GET",  "/api/analisis/historial",    "Devuelve los análisis pasados (perfil + fecha).",           "Historial / Perfil"),
    ]
    for i, row_data in enumerate(ep_data, start=1):
        bg = "F8FAFC" if i % 2 == 0 else None
        add_table_data_row(ep_table, i, row_data, bg_hex=bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading(doc, "5.1 Payload de Análisis Financiero (Ejemplo Real)", level=2, space_before=10)
    add_paragraph(doc, "Entrada al endpoint POST /analisis-financiero:")

    # Código de entrada
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.4)
    code_p.paragraph_format.space_after = Pt(6)
    code_r = code_p.add_run(
        '{\n'
        '  "ingreso_mensual": 4500,\n'
        '  "nivel_endeudamiento": 25,\n'
        '  "frecuencia_ahorro": "Media",\n'
        '  "transacciones": [\n'
        '    { "descripcion": "Supermercado", "valor": 420 },\n'
        '    { "descripcion": "Combustible",  "valor": 300 },\n'
        '    { "descripcion": "Streaming",    "valor": 40  }\n'
        '  ]\n'
        '}'
    )
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(10)
    code_r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    add_paragraph(doc, "Respuesta del endpoint (renderizada dinámicamente en las pantallas de Perfil y Recomendaciones):")

    # Código de salida
    code_p2 = doc.add_paragraph()
    code_p2.paragraph_format.left_indent = Inches(0.4)
    code_p2.paragraph_format.space_after = Pt(12)
    code_r2 = code_p2.add_run(
        '{\n'
        '  "perfil_financiero": "En observación",\n'
        '  "probabilidad": 0.82,\n'
        '  "resumen_gastos": {\n'
        '    "alimentacion": 420,\n'
        '    "transporte":   300,\n'
        '    "entretenimiento": 40\n'
        '  },\n'
        '  "recomendaciones": [\n'
        '    "Monitorear los gastos recurrentes de entretenimiento",\n'
        '    "Aumentar la reserva financiera mensual"\n'
        '  ]\n'
        '}'
    )
    code_r2.font.name = "Courier New"
    code_r2.font.size = Pt(10)
    code_r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 6. DISEÑO VISUAL Y LINEAMIENTOS DE UI/UX
    # ════════════════════════════════════════════════
    add_heading(doc, "6. Lineamientos de Diseño Visual (UI/UX Premium)", level=1)

    add_paragraph(doc,
        "La interfaz de Finance AI está diseñada bajo los principios de diseño fintech moderno: "
        "modo oscuro predeterminado, glassmorphism, colores semánticos adaptativos y micro-animaciones "
        "de 200ms para transmitir interactividad y confianza al usuario."
    )

    add_heading(doc, "6.1 Paleta de Colores Semánticos", level=2, space_before=10)
    color_table = doc.add_table(rows=6, cols=3)
    color_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    color_table.style = "Table Grid"
    add_table_header_row(color_table, ["Elemento / Estado", "Color HEX", "Uso"])
    color_data = [
        ("Fondo principal (dark mode)",  "#0F172A (Slate 900)",   "Background general de la app."),
        ("Tarjetas / Paneles",           "rgba(30,41,59,0.7)",    "Glassmorphism con backdrop-filter: blur(12px)."),
        ("Acento principal (botones)",   "#6366F1 → #3B82F6",     "Gradiente violeta-azul en CTAs y elementos destacados."),
        ("Perfil SALUDABLE",             "#10B981 (Emerald 500)", "Badges, banners y bordes de tarjetas de perfil positivo."),
        ("Perfil EN RIESGO",             "#EF4444 (Red 500)",     "Alertas urgentes, indicadores de riesgo crítico."),
    ]
    for i, row_data in enumerate(color_data, start=1):
        bg = "F8FAFC" if i % 2 == 0 else None
        add_table_data_row(color_table, i, row_data, bg_hex=bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading(doc, "6.2 Tipografía", level=2, space_before=10)
    add_bullet(doc, "Inter (Google Fonts): fuente principal para cuerpo de texto, tablas y formularios.")
    add_bullet(doc, "Plus Jakarta Sans: titulares, KPIs y valores numéricos grandes.")
    add_bullet(doc, "Courier New: bloques de código y valores JSON en la documentación técnica.")

    add_heading(doc, "6.3 Micro-animaciones", level=2, space_before=10)
    add_bullet(doc, "Tarjetas KPI: efecto hover scale(1.02) con transition: 0.2s cubic-bezier(0.4, 0, 0.2, 1).")
    add_bullet(doc, "Gráficos Chart.js: animación de entrada con duration: 800ms y easing: 'easeInOutQuart'.")
    add_bullet(doc, "Semáforo de Perfil: pulso suave (CSS @keyframes pulse) en el badge de estado financiero.")
    add_bullet(doc, "Carga de CSV: barra de progreso animada durante el procesamiento del archivo.")

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 7. CASOS DE USO DETALLADOS
    # ════════════════════════════════════════════════
    add_heading(doc, "7. Casos de Uso del Frontend", level=1)

    add_heading(doc, "CU-01: Registro y Configuración Inicial", level=2, space_before=10)
    add_bullet(doc, "Actor: Usuario nuevo que ingresa por primera vez al sistema.")
    add_bullet(doc, "Flujo: Accede a /register → completa datos → redirige a /onboarding → ingresa ingreso, endeudamiento y frecuencia de ahorro → redirige a /dashboard con perfil base calculado.")
    add_bullet(doc, "Resultado: El usuario tiene un perfil financiero base configurado y puede comenzar a operar en el sistema.")

    add_heading(doc, "CU-02: Clasificación Manual de un Gasto", level=2, space_before=10)
    add_bullet(doc, "Actor: Usuario autenticado que desea registrar un gasto.")
    add_bullet(doc, "Flujo: Va a /transacciones → ingresa 'Cena restaurante' y '$80' → el sistema envía al backend → el modelo clasifica como 'Ocio' → la fila aparece en la tabla con el ícono 🎬 y badge verde-azulado.")
    add_bullet(doc, "Resultado: La transacción queda categorizada y el Dashboard actualiza automáticamente sus gráficos.")

    add_heading(doc, "CU-03: Carga Masiva de Transacciones (CSV)", level=2, space_before=10)
    add_bullet(doc, "Actor: Usuario con extracto bancario mensual en CSV.")
    add_bullet(doc, "Flujo: Arrastra el CSV sobre el componente Drop Zone → el frontend parsea y muestra previsualización de 50 filas → usuario confirma → se envían en lote al backend → tabla se actualiza con todas las categorías clasificadas automáticamente.")
    add_bullet(doc, "Resultado: 50 transacciones clasificadas en segundos sin ingreso manual.")

    add_heading(doc, "CU-04: Simulación de Escenario de Ahorro", level=2, space_before=10)
    add_bullet(doc, "Actor: Usuario que quiere saber qué pasaría si reduce sus gastos de ocio.")
    add_bullet(doc, "Flujo: Va a /simulador → agrega transacciones hipotéticas → el panel derecho muestra el nuevo perfil proyectado ('Saludable' en vez de 'En Observación') → usuario decide si confirma o descarta el escenario.")
    add_bullet(doc, "Resultado: El usuario toma decisiones financieras informadas antes de comprometerse con un gasto.")

    add_heading(doc, "CU-05: Consulta de Recomendaciones Personalizadas", level=2, space_before=10)
    add_bullet(doc, "Actor: Usuario que quiere mejorar su salud financiera.")
    add_bullet(doc, "Flujo: Accede a /recomendaciones → el sistema muestra tarjetas generadas por la IA basadas en su historial de gastos del último mes → el usuario puede marcar cada recomendación como 'Aplicada'.")
    add_bullet(doc, "Resultado: Consejo accionable específico para el comportamiento real del usuario.")

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 8. ARQUITECTURA ANGULAR (ESTRUCTURA DE CÓDIGO)
    # ════════════════════════════════════════════════
    add_heading(doc, "8. Arquitectura Angular Interna (Estructura del Código)", level=1)

    add_paragraph(doc,
        "La aplicación Angular sigue la arquitectura Feature-Oriented con separación clara de capas. "
        "Todos los componentes son Standalone (Angular v17+) y el estado global se gestiona "
        "exclusivamente con Signals sin librerías externas de estado:"
    )

    # Código de estructura
    code_struct = doc.add_paragraph()
    code_struct.paragraph_format.left_indent = Inches(0.4)
    code_struct.paragraph_format.space_after = Pt(12)
    code_r_struct = code_struct.add_run(
        "src/app/\n"
        "├── core/\n"
        "│   ├── services/\n"
        "│   │   ├── finance.service.ts      # Lógica de negocio + Signals globales\n"
        "│   │   └── auth.service.ts         # Gestión del token JWT\n"
        "│   ├── guards/\n"
        "│   │   └── auth.guard.ts           # Protege rutas privadas\n"
        "│   └── interceptors/\n"
        "│       └── auth.interceptor.ts     # Adjunta JWT a cada petición HTTP\n"
        "├── shared/\n"
        "│   └── components/\n"
        "│       ├── kpi-card/               # Tarjeta de indicador reutilizable\n"
        "│       ├── finance-chart/          # Wrapper de Chart.js\n"
        "│       ├── csv-upload/             # Componente Drag & Drop CSV\n"
        "│       └── profile-badge/          # Semáforo de perfil financiero\n"
        "├── features/\n"
        "│   ├── auth/                       # Login + Register\n"
        "│   ├── onboarding/                 # Stepper de configuración inicial\n"
        "│   ├── dashboard/                  # Panel principal\n"
        "│   ├── transactions/               # Tabla + formulario + CSV\n"
        "│   ├── simulator/                  # Simulador de escenarios\n"
        "│   ├── profile/                    # Diagnóstico de perfil\n"
        "│   ├── recommendations/            # Tarjetas de consejos\n"
        "│   ├── history/                    # Gráficos de evolución\n"
        "│   └── settings/                   # Configuración de cuenta\n"
        "├── app.config.ts                   # Proveedores: HttpClient, Router\n"
        "├── app.routes.ts                   # Rutas con Lazy Loading\n"
        "└── app.component.ts                # Shell principal (sidebar + router-outlet)"
    )
    code_r_struct.font.name = "Courier New"
    code_r_struct.font.size = Pt(9)
    code_r_struct.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_page_break()

    # ════════════════════════════════════════════════
    # 9. CONCLUSIÓN Y DIFERENCIACIÓN
    # ════════════════════════════════════════════════
    add_heading(doc, "9. Conclusión y Factores Diferenciadores del MVP", level=1)

    add_paragraph(doc,
        "Finance AI no es simplemente un clasificador de gastos: es un sistema completo de inteligencia "
        "financiera personal. Los factores que lo diferencian y aseguran su destacada calificación en el "
        "Hackathon son:"
    )
    add_bullet(doc, "Sistema completo de 11 pantallas con flujo de usuario real (Login → Onboarding → Dashboard → Análisis → Recomendaciones).", bold_prefix="Completitud: ")
    add_bullet(doc, "Simulador de escenarios hipotéticos que permite al usuario experimentar sin riesgo. Ningún MVP básico del hackathon lo tendrá.", bold_prefix="Diferenciador único: ")
    add_bullet(doc, "Actualización reactiva en tiempo real de todos los indicadores sin recarga de página, gracias a Signals de Angular.", bold_prefix="Reactividad: ")
    add_bullet(doc, "Diseño visual premium tipo fintech moderno: modo oscuro, glassmorphism, animaciones fluidas y colores semánticos adaptativos.", bold_prefix="UI/UX Premium: ")
    add_bullet(doc, "Integración completa con OCI (Object Storage de modelos + Base de Datos + Compute) y arquitectura multicapa desacoplada.", bold_prefix="Integración OCI real: ")

    # ─── GUARDAR ─────────────────────────────────────────────────────────────
    output_path = (
        r"c:\Users\Usuario\Downloads\Proyectos\oracleHackathon"
        r"\g9-latam-team-05\documentacion\documentacionFrontend"
        r"\FinanceAI_Especificacion_Completa_Frontend.docx"
    )
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nDocumento generado exitosamente en:\n{output_path}")

if __name__ == "__main__":
    create_document()
